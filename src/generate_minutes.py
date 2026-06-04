"""文字起こしテキストをGem指示書ルールに従ってWordファイルに変換する"""
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from parse_pdf import MeetingInfo


def to_zenkaku(text: str) -> str:
    """発言内容の半角英数字・記号を全角に変換"""
    result = []
    for ch in text:
        code = ord(ch)
        # 半角英数字・記号（U+0021〜U+007E）→ 全角（U+FF01〜U+FF5E）
        if 0x21 <= code <= 0x7E:
            result.append(chr(code + 0xFEE0))
        # 半角スペース → 全角スペース
        elif ch == " ":
            result.append("　")
        else:
            result.append(ch)
    return "".join(result)


def _is_speaker_line(line: str) -> bool:
    return line.startswith("〇") or line.startswith("○")


def _is_audience_line(line: str) -> bool:
    return line.startswith("（") or line.startswith("(")


def _is_schedule_header(line: str) -> bool:
    return bool(re.match(r"日程第\s*\d+", line))


def _is_time_line(line: str) -> bool:
    return bool(re.match(r"午[前後]\d+時\d*分?(休憩|再開)", line))


def normalize_speaker_line(line: str) -> str:
    """〇役職（氏名君）　内容 の形式に正規化"""
    line = line.replace("○", "〇")
    # 発言部分の半角→全角
    m = re.match(r"(〇[^　\s]+(?:（[^）]+）)?)\s*(.*)", line, re.DOTALL)
    if m:
        prefix = m.group(1)
        content = to_zenkaku(m.group(2).strip())
        return f"{prefix}　{content}"
    return line


def split_into_paragraphs(raw_text: str) -> list[str]:
    """文字起こしテキストを発言単位の段落に分割"""
    lines = raw_text.split("\n")
    paragraphs: list[str] = []
    current: list[str] = []

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if _is_speaker_line(line) or _is_audience_line(line) or _is_schedule_header(line) or _is_time_line(line):
            if current:
                paragraphs.append("\n".join(current))
                current = []
            current.append(line)
        else:
            # 同一人物の継続発言 → 全角スペースで継続
            current.append(line)

    if current:
        paragraphs.append("\n".join(current))

    return paragraphs


def format_paragraph(raw: str) -> list[str]:
    """1段落を出力行リストに変換（Gem指示書ルール適用）"""
    lines_in = raw.split("\n")
    first = lines_in[0].strip()
    rest = lines_in[1:]

    if _is_speaker_line(first):
        # 発言行：〇役職（氏名君）　内容（改行後は全角スペース）
        result_line = normalize_speaker_line(first)
        # 同一人物継続行を結合
        for r in rest:
            r = r.strip()
            if r:
                result_line += "\n　" + to_zenkaku(r)
        return [result_line]

    elif _is_audience_line(first):
        # 傍聴発言はそのまま
        return [first]

    elif _is_schedule_header(first):
        return [first]

    elif _is_time_line(first):
        return [first]

    else:
        lines_out = [first]
        for r in rest:
            r = r.strip()
            if r:
                lines_out.append("　" + r)
        return lines_out


def build_header(info: MeetingInfo) -> list[tuple[str, str]]:
    """文書ヘッダー部分を構築 (text, style) のリスト"""
    rows: list[tuple[str, str]] = []

    # タイトル
    meeting_label = info.title or "椎葉村議会会議録"
    rows.append((f"{meeting_label}（第１日）", "Title"))

    # 日付
    rows.append((info.date or "", "Normal"))

    # 議事日程
    rows.append(("", "Normal"))
    rows.append(("議事日程（第１号）", "Normal"))
    rows.append((f"{info.date}　{info.time}開会" if info.time else info.date, "Normal"))
    rows.append(("", "Normal"))

    for item in info.items:
        rows.append((f"日程第{item.number}　{item.title}", "Normal"))

    rows.append(("", "Normal"))
    return rows


def generate_word(
    transcribed_text: str,
    meeting_info: MeetingInfo,
    output_path: str,
) -> None:
    doc = Document()

    # フォント設定
    style = doc.styles["Normal"]
    style.font.name = "ＭＳ 明朝"
    style.font.size = Pt(10.5)

    # ヘッダー追加
    for text, style_name in build_header(meeting_info):
        p = doc.add_paragraph(text, style=style_name)
        if style_name == "Title":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 本文：文字起こしを段落に分割して追加
    paragraphs = split_into_paragraphs(transcribed_text)
    for para_raw in paragraphs:
        formatted_lines = format_paragraph(para_raw)
        for line in formatted_lines:
            # 改行を含む行は複数パラグラフに分割
            sub_lines = line.split("\n")
            first_sub = True
            for sub in sub_lines:
                if first_sub:
                    doc.add_paragraph(sub, style="Normal")
                    first_sub = False
                else:
                    doc.add_paragraph(sub, style="Normal")

    doc.save(output_path)
    print(f"保存しました: {output_path}")
